from __future__ import annotations

from dataclasses import replace
from datetime import datetime, timezone
import math
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
import uuid

from .chunker import chunk_paragraphs
from .models import Document, DocumentStatus, Paragraph, Section
from .parser import DocumentParser
from .store import DocumentStore


class DocumentService:
    def __init__(self, store: DocumentStore, chunk_size: int, chunk_overlap: int) -> None:
        self._store = store
        self._chunk_size = chunk_size
        self._chunk_overlap = chunk_overlap
        self._parser = DocumentParser()

    def create_document(
        self,
        filename: str,
        content: bytes,
        title: Optional[str] = None,
        source: Optional[str] = None,
        meta: Optional[Dict[str, Any]] = None,
        max_upload_mb: int = 20,
    ) -> Document:
        ext = Path(filename).suffix.lower()
        if ext not in {".pdf", ".docx", ".md", ".markdown", ".txt"}:
            raise ValueError("unsupported file type")
        max_bytes = max_upload_mb * 1024 * 1024
        if len(content) > max_bytes:
            raise ValueError("file too large")

        meta_value = dict(meta or {})
        size_bytes = len(content)
        meta_value.setdefault("size_bytes", size_bytes)
        meta_value.setdefault("size_label", _format_bytes(size_bytes))
        meta_value.setdefault("file_name", filename)
        meta_value.setdefault("version", meta_value.get("version") or "v1.0.0")
        meta_value.setdefault("pages", int(meta_value.get("pages", 0) or 0))
        meta_value.setdefault("chunks", int(meta_value.get("chunks", 0) or 0))
        tags = meta_value.get("tags")
        if isinstance(tags, str):
            tags = [item.strip() for item in tags.split(",") if item.strip()]
        if not tags:
            tags = [ext.lstrip(".") or "doc"]
        meta_value["tags"] = tags

        doc_id = str(uuid.uuid4())
        title_value = title or Path(filename).stem
        source_value = source or filename
        document = Document(
            doc_id=doc_id,
            title=title_value,
            source=source_value,
            status=DocumentStatus.RAW,
            meta=meta_value,
        )
        self._store.save_raw(doc_id, filename, content)
        self._store.save_document(document)
        return document

    def process_document(self, doc_id: str) -> Tuple[List[Section], List[Paragraph]]:
        raw_path = self._store.get_raw_path(doc_id)
        if raw_path is None:
            raise FileNotFoundError("raw file not found")

        try:
            try:
                sections, paragraphs = self._parser.parse(doc_id, raw_path)
            except Exception:
                text = self._fallback_text(raw_path)
                sections, paragraphs = self._parser.parse_plain_text(doc_id, text)
        except Exception as exc:  # noqa: BLE001
            self._set_error(doc_id, exc)
            raise

        if not paragraphs:
            sections, paragraphs = self._parse_empty(doc_id)

        self._store.save_parsed(doc_id, sections, paragraphs)
        self._update_status(doc_id, DocumentStatus.PARSED)

        chunks = chunk_paragraphs(
            doc_id,
            sections,
            paragraphs,
            self._chunk_size,
            self._chunk_overlap,
        )
        self._store.save_chunks(doc_id, chunks)
        self._update_status(doc_id, DocumentStatus.CHUNKED)
        pages = self._estimate_pages(raw_path, paragraphs)
        self._update_meta(
            doc_id,
            {
                "pages": pages,
                "chunks": len(chunks),
            },
        )

        return sections, paragraphs

    def list_documents(self) -> List[Document]:
        docs = self._store.list_documents()
        docs.sort(key=lambda doc: doc.created_at)
        return docs

    def get_document(self, doc_id: str) -> Optional[Dict[str, Any]]:
        doc = self._store.load_document(doc_id)
        if doc is None:
            return None
        parsed = self._store.load_parsed(doc_id)
        chunks = self._store.load_chunks(doc_id)
        return {
            "document": doc.to_dict(),
            "sections": len(parsed["sections"]) if parsed else 0,
            "paragraphs": len(parsed["paragraphs"]) if parsed else 0,
            "chunks": len(chunks) if chunks else 0,
        }

    def get_tree(self, doc_id: str) -> Optional[List[Dict[str, Any]]]:
        parsed = self._store.load_parsed(doc_id)
        if parsed is None:
            return None
        return build_tree(parsed["sections"])

    def _update_status(self, doc_id: str, status: DocumentStatus) -> None:
        doc = self._store.load_document(doc_id)
        if doc is None:
            return
        updated = replace(doc, status=status, updated_at=datetime.now(timezone.utc))
        self._store.save_document(updated)

    def _update_meta(self, doc_id: str, updates: Dict[str, Any]) -> None:
        doc = self._store.load_document(doc_id)
        if doc is None:
            return
        meta = dict(doc.meta or {})
        meta.update(updates)
        updated = replace(doc, meta=meta, updated_at=datetime.now(timezone.utc))
        self._store.save_document(updated)

    def _set_error(self, doc_id: str, exc: Exception) -> None:
        doc = self._store.load_document(doc_id)
        if doc is None:
            return
        meta = dict(doc.meta or {})
        meta["error"] = f"{type(exc).__name__}: {exc}"
        updated = replace(
            doc,
            status=DocumentStatus.ERROR,
            meta=meta,
            updated_at=datetime.now(timezone.utc),
        )
        self._store.save_document(updated)

    def _fallback_text(self, path: Path) -> str:
        raw = path.read_bytes()
        try:
            return raw.decode("utf-8")
        except UnicodeDecodeError:
            return raw.decode("utf-8", errors="ignore")

    def _parse_empty(self, doc_id: str) -> Tuple[List[Section], List[Paragraph]]:
        root = Section(
            section_id=str(uuid.uuid4()),
            doc_id=doc_id,
            level=0,
            title="root",
            path="root",
            parent_id=None,
            order=0,
        )
        return [root], []

    def delete_document(self, doc_id: str) -> bool:
        return self._store.delete_document(doc_id)

    def archive_document(
        self,
        doc_id: str,
        restore: bool = False,
        archive_path: Optional[str] = None,
    ) -> Optional[Document]:
        doc = self._store.load_document(doc_id)
        if doc is None:
            return None
        meta = dict(doc.meta or {})
        if restore:
            previous = meta.pop("archived_from", None)
            if previous:
                try:
                    next_status = DocumentStatus(previous)
                except ValueError:
                    next_status = DocumentStatus.RAW
            else:
                next_status = DocumentStatus.RAW
        else:
            if archive_path is not None:
                cleaned = _normalize_archive_path(archive_path)
                if cleaned:
                    meta["archive_path"] = cleaned
                else:
                    meta.pop("archive_path", None)
            if doc.status != DocumentStatus.ARCHIVED:
                meta["archived_from"] = doc.status.value
            next_status = DocumentStatus.ARCHIVED
        updated = replace(
            doc,
            status=next_status,
            meta=meta,
            updated_at=datetime.now(timezone.utc),
        )
        self._store.save_document(updated)
        return updated

    def _estimate_pages(self, path: Path, paragraphs: List[Paragraph]) -> int:
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            try:
                from PyPDF2 import PdfReader

                reader = PdfReader(str(path))
                return max(1, len(reader.pages))
            except Exception:  # noqa: BLE001
                pass
        if suffix == ".docx":
            try:
                from docx import Document as DocxDocument

                doc = DocxDocument(str(path))
                return max(1, int(math.ceil(len(doc.paragraphs) / 30)))
            except Exception:  # noqa: BLE001
                pass
        if paragraphs:
            return max(1, int(math.ceil(len(paragraphs) / 12)))
        text = self._fallback_text(path)
        return max(1, int(math.ceil(len(text) / 1800)))


def _format_bytes(size: int) -> str:
    if size < 1024:
        return f"{size} B"
    if size < 1024 * 1024:
        return f"{size / 1024:.1f} KB"
    return f"{size / (1024 * 1024):.1f} MB"


def _normalize_archive_path(value: Optional[str]) -> Optional[str]:
    if value is None:
        return None
    trimmed = value.strip()
    if not trimmed:
        return ""
    if trimmed.startswith(("/", "\\")):
        raise ValueError("invalid archive path")
    normalized = trimmed.replace("\\", "/")
    while "//" in normalized:
        normalized = normalized.replace("//", "/")
    normalized = normalized.strip("/")
    if not normalized:
        return ""
    segments = [segment.strip() for segment in normalized.split("/") if segment.strip()]
    for segment in segments:
        if segment in {".", ".."}:
            raise ValueError("invalid archive path")
        if ":" in segment:
            raise ValueError("invalid archive path")
    return "/".join(segments)


def build_tree(sections: List[Section]) -> List[Dict[str, Any]]:
    nodes: Dict[str, Dict[str, Any]] = {}
    roots: List[Dict[str, Any]] = []

    for section in sections:
        nodes[section.section_id] = {
            "id": section.section_id,
            "title": section.title,
            "level": section.level,
            "path": section.path,
            "children": [],
        }

    for section in sections:
        node = nodes[section.section_id]
        if section.parent_id and section.parent_id in nodes:
            nodes[section.parent_id]["children"].append(node)
        else:
            roots.append(node)

    return roots
